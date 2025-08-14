import streamlit as st
import random
import pdfplumber
import re
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
from datetime import datetime

st.title("🏢 Insurance Report Generator")

# Hide Streamlit menu/footer
st.markdown("""
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header > div:nth-child(1) {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# --------------------- INPUTS ---------------------
address = st.text_input("Building Address", "NJ, Newark - 3 Gateway Center")
currency = st.selectbox("Currency", ["CAD", "USD"])
sqft = st.number_input("Building Square Footage", min_value=0.0, value=20000.0)

st.markdown("---")
st.subheader("Optional manual overrides (if OSM data is missing or unavailable)")
multi_tenanted_input = st.selectbox("Is the building multi-tenanted?", ["Unknown", "Yes", "No"], index=0)
building_age_input = st.number_input("Approximate Building Age (years)", min_value=0, max_value=200, value=0)
num_floors_input = st.number_input("Total Floors (excl. basement)", min_value=0, max_value=100, value=0)

st.markdown("---")
st.subheader("📄 Upload Insurance Report PDF file")
pdf_file = st.file_uploader("Upload PDF", type=["pdf"])

DEFAULT_OCR = 0.20  # fallback Occupancy Cost Ratio

# --------------------- LOAD GLOBAL PRICING ---------------------
@st.cache_data
def load_global_pricing():
    file_path = "Global Pricing.xlsx"  # Must exist in repo
    market_rent_df = pd.read_excel(file_path, sheet_name="Market Rent")
    usa_df = pd.read_excel(file_path, sheet_name="USA")
    canada_df = pd.read_excel(file_path, sheet_name="Canada")
    return market_rent_df, usa_df, canada_df

market_rent_df, usa_df, canada_df = load_global_pricing()

# --------------------- HELPER FUNCTIONS ---------------------
def get_address_coords(address):
    # Replace with real geocoding if needed
    return (40.7357, -74.1724)  # Example: Newark, NJ

def get_market_rent_from_address(address):
    addr_coords = get_address_coords(address)
    combined_df = pd.concat([usa_df, canada_df], ignore_index=True)
    combined_df["distance"] = ((combined_df["Latitude"] - addr_coords[0])**2 +
                               (combined_df["Longitude"] - addr_coords[1])**2)**0.5
    nearest_centres = combined_df.nsmallest(3, "distance")["Centre #"].tolist()
    rent_values = market_rent_df[market_rent_df["Centre #"].isin(nearest_centres)]["Market Rate"].tolist()
    if rent_values:
        return sum(rent_values) / len(rent_values)
    return 50.0  # fallback

def extract_gross_area_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                match = re.search(r"Gross Area sqft\s*[:\s]\s*([\d,\.]+)", text, re.IGNORECASE)
                if match:
                    try:
                        return float(match.group(1).replace(',', ''))
                    except:
                        continue
    return None

def extract_value_flexible(table, key_phrase, target_number_index=1):
    key_phrase = key_phrase.lower()
    for row in table:
        row_text = " ".join(str(cell).lower() if cell else "" for cell in row)
        if key_phrase in row_text:
            numbers_str = re.findall(r"[-+]?\d*\.\d+|\d+", row_text.replace(',', ''))
            if len(numbers_str) > target_number_index:
                try:
                    return float(numbers_str[target_number_index])
                except:
                    return None
    return None

def extract_number_next_to_phrase(table, phrase):
    phrase = phrase.lower()
    for row in table:
        for idx, cell in enumerate(row):
            if cell and phrase in str(cell).lower():
                if idx + 1 < len(row):
                    val_str = row[idx + 1]
                    if val_str:
                        val_clean = re.sub(r"[^0-9.\-]", "", str(val_str).replace(',', ''))
                        try:
                            return float(val_clean)
                        except:
                            return None
    return None

def extract_from_pdf(pdf_file):
    payroll = None
    turnover = None
    headline_rent = None
    rentable_area = None
    building_sqft_pdf = extract_gross_area_from_pdf(pdf_file)
    with pdfplumber.open(pdf_file) as pdf:
        all_tables = []
        for page in pdf.pages:
            tables = page.extract_tables()
            if tables:
                all_tables.extend(tables)

        for table in all_tables:
            if payroll is None:
                raw_payroll = extract_value_flexible(table, "staff costs", target_number_index=3)
                if raw_payroll is not None:
                    payroll = raw_payroll * 1000
            if turnover is None:
                raw_turnover = extract_value_flexible(table, "gross revenue", target_number_index=1)
                if raw_turnover is not None:
                    turnover = raw_turnover * 1000

        for table in all_tables:
            if headline_rent is None:
                headline_rent = extract_number_next_to_phrase(table, "headline rent (as reviewed by partner) usd psft p.a.")
            if rentable_area is None:
                rentable_area = extract_number_next_to_phrase(table, "rentable area sqft")

    rental = None
    if headline_rent is not None and rentable_area is not None:
        rental = headline_rent * rentable_area
    return payroll, rental, turnover, building_sqft_pdf

# --------------------- EXTRACT PDF DATA ---------------------
if pdf_file is not None:
    extracted_payroll, extracted_rental, extracted_turnover, pdf_sqft = extract_from_pdf(pdf_file)
    if pdf_sqft:
        sqft = pdf_sqft
else:
    extracted_payroll, extracted_rental, extracted_turnover = None, None, None

# --------------------- UPDATE MARKET RENT ---------------------
market_rent = get_market_rent_from_address(address)

# --------------------- GENERATE REPORT ---------------------
if st.button("Generate Report") and address and sqft > 0 and market_rent > 0:

    multi_tenanted = multi_tenanted_input if multi_tenanted_input != "Unknown" else "Yes"
    building_age = building_age_input if building_age_input > 0 else random.randint(20,50)
    current_year = datetime.now().year
    built_year = current_year - building_age
    num_floors = num_floors_input if num_floors_input > 0 else max(1, int(sqft // 10000))

    if extracted_payroll is None:
        if sqft < 10000:
            fte = 0.5
            payroll = 50000
        elif sqft < 15000:
            fte = 1.0
            payroll = 65000
        elif sqft < 20000:
            fte = 1.5
            payroll = 110000
        else:
            fte = 2.0
            payroll = 110000
    else:
        payroll = extracted_payroll
        fte = round(payroll / 55000,1)

    rental_estimate = extracted_rental if extracted_rental is not None else sqft * market_rent
    annual_turnover = extracted_turnover if extracted_turnover is not None else (rental_estimate / DEFAULT_OCR)
    gross_profit_calc = annual_turnover - rental_estimate if annual_turnover and rental_estimate else None

    st.subheader("🏗 Building Information")
    st.write(f"**Address:** {address}")
    st.write(f"**Multi-tenanted:** {multi_tenanted}")
    st.write(f"**Approximate Age:** {building_age} years (built in {built_year})")
    st.write(f"**Total Floors:** {num_floors}")

    st.subheader("Employment Estimate")
    st.write(f"**Estimated FTEs:** {fte}")
    st.write(f"**Estimated Annual Payroll:** {currency} {payroll:,.2f}")

    st.subheader("Forecasted Financials")
    st.write(f"**Rental (Budget/Estimate - Next Year):** {currency} {rental_estimate:,.2f}")
    st.write(f"**Annual Turnover (Forecast):** {currency} {annual_turnover:,.2f}")
    st.write(f"**Annual Gross Profit (calculated):** {currency} {gross_profit_calc:,.2f}")

    try:
        doc = Document("Insurance Template.docx")

        for paragraph in doc.paragraphs:
            if "All values are in" in paragraph.text:
                paragraph.text = f"Please see the answers in blue below. All values are in {currency} and sq ft."
            if "Is building multi- tenanted" in paragraph.text:
                paragraph.add_run(f" {multi_tenanted}").font.color.rgb = RGBColor(0,0,255)
            if "Approximate age of the building" in paragraph.text:
                paragraph.add_run(f" {building_age} years (built in {built_year})").font.color.rgb = RGBColor(0,0,255)
            if "Total number of floors" in paragraph.text:
                paragraph.add_run(f" {num_floors}").font.color.rgb = RGBColor(0,0,255)
            if "Number of employees will be employed" in paragraph.text:
                paragraph.add_run(f" {fte}").font.color.rgb = RGBColor(0,0,255)

        table_values = [annual_turnover, gross_profit_calc, rental_estimate, payroll]
        for table in doc.tables:
            val_idx = 0
            for row in table.rows:
                for cell in row.cells:
                    if "$" in cell.text and val_idx < len(table_values):
                        cell.text = f"{currency} {table_values[val_idx]:,.2f}"
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0,0,255)
                        val_idx += 1

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        st.download_button(
            label="📄 Download Insurance Report",
            data=output,
            file_name=f"Insurance_Report_{address.replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"Error generating Word document: {e}")
